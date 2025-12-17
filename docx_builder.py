from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from constants import TARGET_WIDTH, TARGET_HEIGHT
from pdf_utils import pdf_to_images
from masking_visual import apply_visual_mask  # somente TR


# =================================================
# FUNÇÃO PRINCIPAL
# =================================================
def build_docx(tr_bytes, proposal_files, debug=False):
    """
    Constrói um único DOCX com a seguinte ordem fixa:
    1) Termo de Referência (com mascaramento visual)
    2) Proposta(s) de Preços (sem mascaramento)

    Parâmetros:
    - tr_bytes: bytes do PDF do Termo de Referência
    - proposal_files: lista de bytes dos PDFs de propostas
    - debug: ativa visualização da máscara (uso técnico)
    """

    doc = Document()
    _configure_page(doc)

    # -------------------------------------------------
    # 1️⃣ PROCESSAMENTO DO TERMO DE REFERÊNCIA (TR)
    # -------------------------------------------------
    tr_images = _process_tr(tr_bytes, debug)

    if not tr_images:
    raise RuntimeError(
        "ERRO CRÍTICO: Nenhuma página gerada para o Termo de Referência."
    )

for idx, img in enumerate(tr_images):
    _add_image_to_docx(doc, img)
    doc.add_page_break()


    # -------------------------------------------------
    # 2️⃣ PROCESSAMENTO DAS PROPOSTAS (SEM MÁSCARA)
    # -------------------------------------------------
    for proposal_bytes in proposal_files:
        proposal_images = _process_proposal(proposal_bytes)

        for img in proposal_images:
            doc.add_page_break()
            _add_image_to_docx(doc, img)

    # -------------------------------------------------
    # SAÍDA
    # -------------------------------------------------
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# =================================================
# FUNÇÕES AUXILIARES – TR
# =================================================
def _process_tr(pdf_bytes, debug):
    """
    Processa o Termo de Referência:
    - converte PDF em imagens
    - aplica mascaramento visual SOMENTE nas páginas com tabela
    - ajusta escala final
    """
    images = pdf_to_images(pdf_bytes)
    processed_images = []

    for img in images:
        img = apply_visual_mask(img, debug=debug)
        img = img.resize((TARGET_WIDTH, TARGET_HEIGHT))
        processed_images.append(img)

    return processed_images


# =================================================
# FUNÇÕES AUXILIARES – PROPOSTAS
# =================================================
def _process_proposal(pdf_bytes):
    """
    Processa Proposta de Preços:
    - converte PDF em imagens
    - NÃO aplica qualquer mascaramento
    - ajusta escala final
    """
    images = pdf_to_images(pdf_bytes)
    processed_images = []

    for img in images:
        img = img.resize((TARGET_WIDTH, TARGET_HEIGHT))
        processed_images.append(img)

    return processed_images


# =================================================
# FUNÇÕES AUXILIARES – DOCX
# =================================================
def _add_image_to_docx(doc, image):
    """
    Insere imagem no DOCX:
    - centralizada
    - sem espaçamento adicional
    """
    bio = BytesIO()
    image.save(bio, format="JPEG", quality=85, optimize=True)
    bio.seek(0)

    doc.add_picture(bio, width=Cm(18))
    par = doc.paragraphs[-1]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = 0
    par.paragraph_format.space_after = 0


def _configure_page(doc):
    """
    Configuração fixa de página (SEI):
    """
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.5)
