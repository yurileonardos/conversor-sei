import streamlit as st
from pdf2image import convert_from_bytes
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import zipfile
import pdfplumber
from PIL import ImageDraw

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="SEI Converter ATA - SGB",
    page_icon="📑",
    layout="centered"
)

# --- ESTILO CSS ---
hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            .footer {
                position: fixed;
                left: 0;
                bottom: 0;
                width: 100%;
                background-color: #f1f1f1;
                color: #555;
                text-align: center;
                padding: 10px;
                font-size: 14px;
                z-index: 999;
            }
            .stFileUploader label {
                 font-size: 18px;
                 font-weight: bold;
            }
            .stFileUploader {
                padding: 20px;
                border-radius: 10px;
                border: 2px dashed #cccccc;
            }
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# --- TÍTULO PRINCIPAL ---
st.title("📑 SEI Converter ATA - SGB")

# --- INTRODUÇÃO ---
st.markdown("""
Converta documentos PDF de **TR (Termo de Referência)** e **Proposta de Preços** em imagens otimizadas, 
a fim de inseri-las no documento SEI: **ATA DE REGISTRO DE PREÇOS**.
""")

col1, col2 = st.columns([0.1, 0.9])
with col1:
    try:
        st.image("icone_sei.png", width=40)
    except:
        st.write("🧩")
with col2:
    st.info("""
    Funcionalidade disponível na extensão [**SEI PRO**](https://sei-pro.github.io/sei-pro/), 
    utilizando a ferramenta [**INSERIR CONTEÚDO EXTERNO**](https://sei-pro.github.io/sei-pro/pages/INSERIRDOC.html).
    """)

with st.expander("⚙️ Deseja escolher a pasta onde o arquivo será salvo? Clique aqui."):
    st.markdown("""
    Por segurança, os navegadores salvam automaticamente na pasta "Downloads". 
    Para escolher a pasta a cada download, configure seu navegador (Chrome/Edge):
    1. Vá em **Configurações** > **Downloads**.
    2. Ative: **"Perguntar onde salvar cada arquivo antes de fazer download"**.
    """)

st.write("---")

# --- PASSO 1: UPLOAD ---
st.write("### Passo 1: Upload dos Arquivos")
st.markdown("**Nota:** O sistema aplicará a máscara automática para ocultar preços em Termos de Referência.")

uploaded_files = st.file_uploader(
    "Arraste e solte seus arquivos PDF aqui (ou clique para buscar):", 
    type="pdf", 
    accept_multiple_files=True
)

# --- FUNÇÃO AUXILIAR DE LIMPEZA DE TEXTO ---
def clean_text(text):
    if not text: return ""
    text = text.replace('\n', ' ').replace('\r', ' ')
    return text.lower().strip()

# --- FUNÇÃO DE MASCARAMENTO (v11.0 - SAFETY BRAKE) ---
def apply_masking_v11(image, pdf_page, mask_state):
    """
    Lógica aprimorada para evitar danos colaterais em páginas de texto.
    mask_state: {'mask_x': float, 'last_table_bbox': list, 'table_type': str}
    """
    
    # 1. Tenta achar tabelas reais (com linhas) - Prioridade Alta
    tables_lines = pdf_page.find_tables(table_settings={"vertical_strategy": "lines", "horizontal_strategy": "lines"})
    
    # 2. Tenta achar tabelas implícitas (texto alinhado) - Prioridade Baixa
    tables_text = pdf_page.find_tables(table_settings={"vertical_strategy": "text", "horizontal_strategy": "text"})
    
    # Decide qual usar. Se tivermos linhas, usamos linhas. Se não, usamos texto.
    # Mas marcamos o tipo para saber se mudou o padrão.
    if tables_lines:
        current_tables = tables_lines
        current_type = "lines"
    else:
        current_tables = tables_text
        current_type = "text"

    draw = ImageDraw.Draw(image)
    im_width, im_height = image.size
    scale_x = im_width / pdf_page.width
    scale_y = im_height / pdf_page.height

    # Palavras-chave de Preço (Ativação)
    keywords_price = [
        "preco unit", "preço unit", "valor unit", "vlr. unit", "unitario", "unitário",
        "valor max", "valor estim", "preço estim", "preco estim", "valor ref", 
        "vlr total", "valor total", "preco total", "preço total"
    ]
    
    # Palavras-chave de Cabeçalhos Genéricos (Desativação/Segurança)
    # Se encontrarmos isso SEM preço, desligamos a máscara.
    keywords_generic_header = [
        "item", "descrição", "especificação", "produto", "local", "prazo", 
        "responsável", "etapa", "endereço", "unidade", "catmat", "quantidade"
    ]

    for table in current_tables:
        if not table.rows: continue

        # --- ANÁLISE DO CABEÇALHO DA TABELA ATUAL ---
        has_price_header = False
        has_generic_header = False
        new_mask_x = None

        # Analisa as primeiras linhas buscando pistas
        for row_idx in range(min(3, len(table.rows))):
            row_cells = table.rows[row_idx].cells
            for cell_idx, cell in enumerate(row_cells):
                if not cell: continue
                try:
                    cropped = pdf_page.crop(cell)
                    text_raw = cropped.extract_text()
                    text_clean = clean_text(text_raw)
                    
                    # Checa se é preço
                    if any(k in text_clean for k in keywords_price):
                        has_price_header = True
                        new_mask_x = cell[0]
                    
                    # Checa se é um cabeçalho genérico (mas não preço)
                    if any(k in text_clean for k in keywords_generic_header):
                        has_generic_header = True
                except:
                    pass
            if has_price_header: break

        # --- LÓGICA DE DECISÃO (SAFETY BRAKE) ---
        active_mask_x = None

        if has_price_header:
            # CASO 1: É explicitamente uma tabela de preço.
            # Ativa/Atualiza a máscara
            mask_state['mask_x'] = new_mask_x
            mask_state['last_table_bbox'] = table.bbox
            mask_state['table_type'] = current_type
            active_mask_x = new_mask_x
        
        elif has_generic_header and not has_price_header:
            # CASO 2: É uma tabela nova (ex: Locais de Entrega), mas SEM preço.
            # FREIO DE SEGURANÇA: Desliga a máscara imediatamente.
            mask_state['mask_x'] = None
            mask_state['last_table_bbox'] = None
            mask_state['table_type'] = None
            active_mask_x = None # Garante que não desenha nada
        
        else:
            # CASO 3: Não tem cabeçalho claro (provável continuação ou texto solto)
            # Só aplicamos se a memória estiver ativa E o tipo de tabela for compatível
            if mask_state['mask_x'] is not None:
                
                # Verificação de Tipo: Se a tabela original tinha linhas e essa é "texto",
                # é provável que a tabela acabou e virou parágrafo. Aborta.
                if mask_state['table_type'] == 'lines' and current_type == 'text':
                    mask_state['mask_x'] = None # Desliga para garantir
                    active_mask_x = None
                else:
                    # Verificação Geométrica (Alinhamento)
                    prev_bbox = mask_state['last_table_bbox']
                    curr_bbox = table.bbox
                    if prev_bbox:
                        # Se estiver alinhada à esquerda com tolerância
                        if abs(curr_bbox[0] - prev_bbox[0]) < 40: 
                            active_mask_x = mask_state['mask_x']
                            mask_state['last_table_bbox'] = table.bbox # Atualiza para a próxima
                        else:
                            # Desalinhou muito? Provavelmente outra coisa.
                            mask_state['mask_x'] = None 

        # --- DESENHO DA MÁSCARA ---
        if active_mask_x is not None:
            table_rect = table.bbox
            
            if active_mask_x < table_rect[2]:
                # Máscara Branca
                rect_mask = [
                    active_mask_x * scale_x,       
                    table_rect[1] * scale_y,      
                    table_rect[2] * scale_x + 50, 
                    table_rect[3] * scale_y       
                ]
                draw.rectangle(rect_mask, fill="white", outline=None)

                # Borda Preta
                rect_border = [
                    table_rect[0] * scale_x,
                    table_rect[1] * scale_y,
                    active_mask_x * scale_x,
                    table_rect[3] * scale_y
                ]
                draw.rectangle(rect_border, outline="black", width=2)
            
            # Limpeza de Rodapé (Total)
            try:
                last_row = table.rows[-1]
                first_cell = last_row.cells[0]
                if first_cell:
                    cropped_last = pdf_page.crop(first_cell)
                    last_text = clean_text(cropped_last.extract_text())
                    if "total" in last_text:
                        tops = [c[1] for c in last_row.cells if c]
                        bottoms = [c[3] for c in last_row.cells if c]
                        if tops and bottoms:
                            rect_total = [
                                table.bbox[0] * scale_x,
                                min(tops) * scale_y,
                                active_mask_x * scale_x, 
                                max(bottoms) * scale_y
                            ]
                            draw.rectangle(rect_total, fill="white", outline="black", width=2)
            except:
                pass

    return image, mask_state

# --- FUNÇÃO DE CONVERSÃO ---
def convert_pdf_to_docx(file_bytes):
    try:
        pdf_plumb = pdfplumber.open(BytesIO(file_bytes))
        has_text_layer = True
    except:
        has_text_layer = False
        pdf_plumb = None

    images = convert_from_bytes(file_bytes)
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.5)

    # Estado da Máscara (Reset por arquivo)
    mask_state = {'mask_x': None, 'last_table_bbox': None, 'table_type': None}

    for i, img in enumerate(images):
        if has_text_layer and pdf_plumb and i < len(pdf_plumb.pages):
            img, mask_state = apply_masking_v11(img, pdf_plumb.pages[i], mask_state)
        
        img = img.resize((595, 842)) 
        img_byte_arr = BytesIO()
        img.save(img_byte_arr, format='JPEG', quality=80, optimize=True)
        img_byte_arr.seek(0)

        doc.add_picture(img_byte_arr, width=Cm(18.0))
        
        par = doc.paragraphs[-1]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        
        if i < len(images) - 1:
            doc.add_page_break()
    
    docx_io = BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io

# --- PASSO 2: PROCESSAR ---
if uploaded_files:
    st.write("---")
    st.write("### Passo 2: Converter e Download")
    
    qtd = len(uploaded_files)
    st.caption(f"{qtd} arquivo(s) pronto(s) para conversão.")

    if st.button(f"🚀 Processar Arquivos"):
        with st.spinner('Analisando e protegendo dados...'):
            try:
                processed_files = []
                progress_bar = st.progress(0)
                
                for index, uploaded_file in enumerate(uploaded_files):
                    docx_data = convert_pdf_to_docx(uploaded_file.read())
                    file_name = uploaded_file.name.replace('.pdf', '') + "_SEI_SGB.docx"
                    processed_files.append((file_name, docx_data))
                    progress_bar.progress((index + 1) / qtd)

                st.success("✅ Processamento concluído!")
                
                if len(processed_files) == 1:
                    name, data = processed_files[0]
                    st.download_button(
                        label=f"📥 Salvar {name} no Computador",
                        data=data,
                        file_name=name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    zip_buffer = BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w") as zf:
                        for name, data in processed_files:
                            zf.writestr(name, data.getvalue())
                    zip_buffer.seek(0)
                    st.download_button(
                        label="📥 Salvar Todos (.ZIP) no Computador",
                        data=zip_buffer,
                        file_name="Documentos_SEI_Convertidos.zip",
                        mime="application/zip"
                    )

            except Exception as e:
                st.error(f"Erro: {e}")

# --- GUIA VISUAL ---
st.write("---")
st.subheader("📚 Guia Rápido: Como inserir no SEI")

col1, col2 = st.columns([0.15, 0.85])
with col1:
    try:
        st.image("icone_sei.png", width=50) 
    except:
        st.write("🧩")
with col2:
    st.markdown("""
    *1º Localize o ícone:* No editor do SEI, clique no botão da função *INSERIR CONTEÚDO EXTERNO* (representado pelo ícone ao lado).
    """)

st.write("")

st.markdown("""
*2º Configure a inserção:* Na janela que abrir, faça o upload do arquivo Word gerado aqui.
""")

st.warning("⚠️ *IMPORTANTE:* Certifique-se de deixar todas as caixas de seleção *DESMARCADAS*.")

try:
    st.image("print_sei.png", caption="Exemplo: Deixe as opções desmarcadas.", use_container_width=True)
except:
    pass

# --- RODAPÉ ---
st.markdown('<div class="footer">Developed by Yuri 🚀 | SEI Converter ATA - SGB v11.0 (Safety Fix)</div>', unsafe_allow_html=True)
